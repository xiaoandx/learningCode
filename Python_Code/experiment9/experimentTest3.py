#!/usr/bin/env python
# -*- coding: UTF-8 -*-
"""
 Copyright (c) 2020 WEI.ZHOU. All rights reserved.
 The following code snippets are only used for circulation and cannot be used for business.
 If the code is used, no consent is required, but the author has nothing to do with any problems and
 -consequences.

 In case of code problems, feedback can be made through the following email address.
                         <xiaoandx@gmail.com>

 @Description: word正则匹配AABB模式字符
 @File: experimentTest3.py
 @Author: WEI.ZHOU
 @Date: 2020-12-10 13:46:48
 @Version: V1.0
 @Others:  Running test instructions
"""
import docx
import re


def printResult(strArray=[]):
    """
    匹配结果输出
    :param strArray:
    :return:
    """
    a = 0
    for t in strArray:
        if len(t) == 4:
            if a % 2 == 0:
                print(t, end=" ")
        a += 1


def printDefTxt(file):
    """
    输出docx原文
    :param file: docx文件
    :return:
    """
    print("=" * 40)
    print("字符原文：")
    # 输出原文
    for para in file.paragraphs:
        print(para.text)


def matching(file, pattern):
    """
    正则匹配
    :param file:
    :param pattern:
    :return:
    """
    strArray = []
    # 输出每一段的内容
    for para in file.paragraphs:
        testArr = re.findall(pattern, ''.join(para.text.split("、")))
        if len(testArr):
            for i in testArr:
                for j in i:
                    strArray.append(j)

    print("\n"+"=" * 40)
    print("匹配字符：")
    printResult(strArray)


def main():
    # 正则表达式
    pattern = r'(((.).\3.)|((.)\5(.)\6))'
    # docx路径
    fileUrl = ".\\temp\\test.docx"
    # docx文本对象
    file = docx.Document(fileUrl)
    printDefTxt(file)
    matching(file, pattern)


if __name__ == '__main__':
    main()
